from docx import Document
from docx.shared import Inches


document = Document()

p = document.add_heading('Document title', 0)

p.add_run('bold').bold = True
p.add_run('and some')
p.add_run('italic').italic = True


